import os
import PyPDF2
import docx
import requests
import time
import logging
from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.agents import initialize_agent, Tool
from langchain_community.tools.yahoo_finance_news import YahooFinanceNewsTool
from langchain.agents.agent_types import AgentType
from langchain_openai import AzureChatOpenAI, AzureOpenAIEmbeddings
from langchain_postgres.vectorstores import PGVector
import psycopg2
import openpyxl  # For .xlsx files
import xlrd  # For .xls files
from concurrent.futures import ThreadPoolExecutor
import json  # Add this import at the top of your file
from dotenv import load_dotenv  # Import dotenv to load environment variables
import re  # Add this import at the top of your file

# Load environment variables from .env file
load_dotenv()

# ---------- CONFIGURATION ---------- #
UPLOAD_FOLDER = os.getenv("UPLOAD_FOLDER", "./uploads")
ALLOWED_EXTENSIONS = set(os.getenv("ALLOWED_EXTENSIONS", "pdf,docx,xlsx,xls").split(","))

DB_CONNECTION_STRING_SQLALCHEMY = (
    f"postgresql+psycopg2://{os.getenv('DB_USER')}:{os.getenv('DB_PASSWORD')}@"
    f"{os.getenv('DB_HOST')}:{os.getenv('DB_PORT')}/{os.getenv('DB_NAME')}"
)
DB_CONNECTION_STRING_PSYCOPG2 = (
    f"dbname={os.getenv('DB_NAME')} "
    f"user={os.getenv('DB_USER')} "
    f"password={os.getenv('DB_PASSWORD')} "
    f"host={os.getenv('DB_HOST')} "
    f"port={os.getenv('DB_PORT')} sslmode=require"
)

# Set USER_AGENT from environment variable or use a default value
USER_AGENT = os.getenv(
    "USER_AGENT",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/108.0.0.0 Safari/537.36"
)


try:
    conn = psycopg2.connect(
        "host=dataopspgvectorserver.postgres.database.azure.com "
        "dbname=StockAnalysis user=sentiment password=Admin1234 port=5432 sslmode=require"
    )
    print("Connection successful!")
    conn.close()
except Exception as e:
    print(f"Connection failed: {e}")

# ---------- DATABASE SETUP ---------- #
def create_pgvector_table():
    """Create the required tables for PGVector and analysis results if they don't exist."""
    try:
        conn = psycopg2.connect(DB_CONNECTION_STRING_PSYCOPG2)
        cursor = conn.cursor()

        # Table for embeddings
        create_embedding_table_query = """
        CREATE TABLE IF NOT EXISTS langchain_pg_embedding (
            id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            collection_id UUID NOT NULL,
            embedding VECTOR NOT NULL,
            document TEXT NOT NULL,
            cmetadata JSONB
        );
        """
        cursor.execute(create_embedding_table_query)

        # Table for analysis results
        create_analysis_table_query = """
        CREATE TABLE IF NOT EXISTS analysis_results (
            id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            company_name TEXT NOT NULL,
            ticker TEXT NOT NULL,
            analysis JSONB NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
        """
        cursor.execute(create_analysis_table_query)

        cursor.execute("CREATE EXTENSION IF NOT EXISTS vector;")

        conn.commit()
        cursor.close()
        conn.close()
        logging.info("Tables 'langchain_pg_embedding' and 'analysis_results' are ready.")
    except Exception as e:
        logging.error(f"Error creating tables: {str(e)}")
        raise
# ---------- FLASK APP SETUP ---------- #
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ---------- LOGGING SETUP ---------- #
logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("app.log"),
        logging.StreamHandler()
    ]
)

# ---------- CLIENT OBJECT ---------- #
llm_client = AzureChatOpenAI(
    deployment_name=os.environ["AZURE_CHAT_DEPLOYMENT"],
    azure_endpoint=os.environ["AZURE_CHAT_ENDPOINT"],
    api_key=os.environ["AZURE_CHAT_API_KEY"],
    api_version=os.environ["AZURE_CHAT_API_VERSION"],
    temperature=0
)
 
embedding_client = AzureOpenAIEmbeddings(
    deployment=os.environ["AZURE_EMBEDDING_DEPLOYMENT"],
    azure_endpoint=os.environ["AZURE_EMBEDDING_ENDPOINT"],
    api_key=os.environ["AZURE_EMBEDDING_API_KEY"],
    api_version=os.environ["AZURE_EMBEDDING_API_VERSION"]
)

# ---------- UTILITIES ---------- #
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def load_document(file_path):
    logging.info("Step 0: Loading document and extracting text...")
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    if ext == ".pdf":
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"

    elif ext == ".docx":
        doc = docx.Document(file_path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    elif ext == ".xlsx":
        workbook = openpyxl.load_workbook(file_path)
        for sheet in workbook.sheetnames:
            worksheet = workbook[sheet]
            for row in worksheet.iter_rows(values_only=True):
                text += " ".join([str(cell) for cell in row if cell is not None]) + "\n"

    elif ext == ".xls":
        workbook = xlrd.open_workbook(file_path)
        for sheet in workbook.sheets():
            for row_idx in range(sheet.nrows):
                row = sheet.row(row_idx)
                text += " ".join([str(cell.value) for cell in row if cell.value]) + "\n"

    else:
        raise ValueError("Unsupported file type")

    logging.info("Step 1 Complete: Text extracted.")
    return text

def extract_company_name(text):
    logging.info("Step 2: Extracting company name using GPT-4...")
    prompt = f"""
Given the quarterly business report below, extract the full legal company name (e.g., 'Apple Inc', 'Tesla, Inc'):

{text[:2000]}

Respond with only the company name.
"""
    response = llm_client.invoke(prompt)
    company_name = response.content.strip()
    if not company_name:
        raise ValueError("Could not extract company name")
    logging.info(f"Step 2 Complete: Company name extracted -> {company_name}")
    return company_name

def get_ticker(company_name):
    """
    Fetch the ticker symbol for a company using Yahoo Finance API.

    Args:
        company_name (str): The name of the company to search for.

    Returns:
        str: The ticker symbol of the company, or None if not found.
    """
    logging.info(f"Fetching ticker for company: {company_name}...")
    yfinance = "https://query2.finance.yahoo.com/v1/finance/search"
    params = {"q": company_name, "quotes_count": 1, "country": "United States"}

    retries = 3
    for attempt in range(retries):
        try:
            res = requests.get(url=yfinance, params=params, headers={'User-Agent': USER_AGENT}, timeout=30)
            res.raise_for_status()
            data = res.json()

            if "quotes" in data and len(data["quotes"]) > 0:
                company_code = data['quotes'][0]['symbol']
                logging.info(f"Ticker found for {company_name}: {company_code}")
                return company_code
            else:
                logging.warning(f"No ticker found for {company_name}.")
                return None
        except requests.exceptions.RequestException as e:
            logging.error(f"Attempt {attempt + 1} failed: {e}")
            if attempt < retries - 1:
                time.sleep(2 ** attempt)  # Exponential backoff
            else:
                raise

def fetch_company_news_with_agent(company_name):
    """
    Fetch the latest news about a company using YahooFinanceNewsTool and LangChain agent.
    
    Args:
        company_name (str): The name of the company to search for.
    
    Returns:
        str: A summary of the latest news about the company.
    """
    logging.info(f"Fetching latest news for company: {company_name} using LangChain agent...")

    tools = [YahooFinanceNewsTool()]
    agent_chain = initialize_agent(
        tools,
        llm_client,
        agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
        verbose=True,
    )

    # Use the agent to fetch news
    prompt = f"Get the latest news about {company_name}."
    result = agent_chain.invoke(prompt)

    logging.info(f"Fetched news for {company_name}: {result}")
    return result

def embed_and_store(text, collection_name):
    logging.info("Step 5: Splitting document and storing in pgvector...")
    splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
    chunks = splitter.split_text(text)
    batch_size = 10  # Process 10 chunks at a time

    def process_batch(batch):
        docs = [Document(page_content=chunk) for chunk in batch]
        PGVector.from_documents(
            documents=docs,
            embedding=embedding_client,
            collection_name=collection_name,
            connection=DB_CONNECTION_STRING_SQLALCHEMY
        )

    # Use ThreadPoolExecutor to process batches in parallel
    with ThreadPoolExecutor(max_workers=4) as executor:  # Adjust max_workers based on your system's resources
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            executor.submit(process_batch, batch)

    logging.info("Step 5 Complete: Document stored in pgvector.")

import re  # Add this import at the top of your file

def run_agent_analysis(ticker, collection_name, company_name, document_text):
    logging.info("Step 6: Running agentic analysis using GPT-4...")

    # Fetch the latest news for the company
    news_text = fetch_company_news_with_agent(company_name)

    # Create an instance of PGVector
    pgvector_instance = PGVector(
        embeddings=embedding_client,
        collection_name=collection_name,
        connection=DB_CONNECTION_STRING_SQLALCHEMY
    )

    # Call as_retriever on the instance
    retriever = pgvector_instance.as_retriever()

    tools = [
        Tool(
            name="QuarterlyReportTool",
            func=lambda q: "\n".join([doc.page_content for doc in retriever.get_relevant_documents(q)]),
            description="Returns relevant sections from the quarterly report"
        ),
        Tool(
            name="CompanyNewsTool",
            func=lambda _: news_text,
            description="Returns the latest news about the company"
        )
    ]

    agent = initialize_agent(
        tools,
        llm_client,
        agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
        verbose=True
    )

    final_prompt = f"""
Using the tools available (quarterly report + latest news), analyze the following:
1. What is the overall sentiment (positive, neutral, negative)?
2. Summarize quarterly report highlights for {company_name}.
3. Summarize the latest news and its sentiment.
4. Final recommendation: Buy / Hold / Sell? Justify.
"""
    result = agent.invoke(final_prompt)
    logging.info("Step 6 Complete: Analysis complete.")

    # Clean the result to remove special characters, patterns like \n before numbers, and extra spaces
    def clean_text(text):
        # Remove \n, \r, \t
        text = re.sub(r'[\n\r\t]', ' ', text)
        # Remove \n before all numbers (e.g., \n1 -> 1, \n2 -> 2)
        text = re.sub(r'\s*\n(\d+)', r' \1', text)
        # Replace multiple spaces with a single space
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    if isinstance(result, dict):
        # Clean all string values in the result dictionary
        cleaned_result = {key: clean_text(value) if isinstance(value, str) else value for key, value in result.items()}
    else:
        # If the result is not a dictionary, clean it directly
        cleaned_result = clean_text(result)

    # Store the cleaned analysis result in the database
    try:
        conn = psycopg2.connect(DB_CONNECTION_STRING_PSYCOPG2)
        cursor = conn.cursor()

        # Convert the cleaned result dictionary to a JSON string
        result_json = json.dumps(cleaned_result)

        insert_query = """
        INSERT INTO analysis_results (company_name, ticker, analysis)
        VALUES (%s, %s, %s);
        """
        cursor.execute(insert_query, (company_name, ticker, result_json))
        conn.commit()
        cursor.close()
        conn.close()
        logging.info("Analysis result stored in the database.")
    except Exception as e:
        logging.error(f"Error storing analysis result: {str(e)}")
        raise

    return cleaned_result

# ---------- FLASK API ---------- #
@app.route('/upload', methods=['POST'])
def upload_report():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']

    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({"error": "Invalid file format"}), 400

    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)

    try:
        logging.info("=" * 60)
        logging.info(f"File Uploaded: {filename}")

        # Load the document and extract text
        text = load_document(file_path)

        # Extract the company name from the document
        company_name = extract_company_name(text)

        # Fetch the ticker symbol for the extracted company name
        ticker = get_ticker(company_name)
        if not ticker:
            raise ValueError(f"Could not find ticker for company: {company_name}")

        # Embed and store the document
        embed_and_store(text, ticker)

        # Run analysis with the document and fetched news
        analysis = run_agent_analysis(ticker, ticker, company_name, text)  # Pass `text` as `document_text`
        print(analysis)
        logging.info("Analysis complete.")

        return jsonify({
            "company_name": company_name,
            "ticker": ticker,
            "analysis": analysis
        }), 200

    except Exception as e:
        logging.error(f"ERROR: {str(e)}")
        return jsonify({"error": str(e)}), 500
    
@app.route('/query_analysis', methods=['GET'])
def query_analysis():
    """
    Query analysis results from the database based on ticker.
    """
    ticker = request.args.get('ticker')

    if not ticker:
        return jsonify({"error": "Please provide 'ticker' as a query parameter."}), 400

    try:
        # Connect to the PostgreSQL database
        conn = psycopg2.connect(DB_CONNECTION_STRING_PSYCOPG2)
        cursor = conn.cursor()

        # Build the query to filter by ticker
        query = """
        SELECT company_name, ticker, analysis, created_at 
        FROM analysis_results 
        WHERE ticker = %s
        """
        params = [ticker]

        # Execute the query
        cursor.execute(query, params)
        results = cursor.fetchall()

        # Format the results
        analysis_list = []
        for row in results:
            analysis_list.append({
                "company_name": row[0],
                "ticker": row[1],
                "analysis": row[2],
                "created_at": row[3].isoformat()
            })

        cursor.close()
        conn.close()

        if not analysis_list:
            return jsonify({"message": "No analysis results found for the given ticker."}), 404

        return jsonify({"results": analysis_list}), 200

    except Exception as e:
        logging.error(f"Error querying analysis results: {str(e)}")
        return jsonify({"error": str(e)}), 500
    
# ---------- MAIN ---------- #
if __name__ == '__main__':
    create_pgvector_table()
    app.run(debug=True)